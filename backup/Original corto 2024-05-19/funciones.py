import plotly.graph_objs as go
import pandas as pd
import numpy as np
import streamlit as st

import docx

from snowflake_utils import sf_check_snowflake_connection, st_query_to_snowflake_and_return_dataframe
from snowflake_config import sf_config

# Cargar dataframes desde Snowflake. Como están en funciones, sólo se cargan una vez y optimiza el consumo. 
# Si se dejan en el app.py cada vez que se consulta un departamento nuevo, se hace la conexión a Snowflake. 
expected_types = {'Cod. Municipio': str}
df_general = st_query_to_snowflake_and_return_dataframe("""
    SELECT * FROM TABLA_BASE_MUNICIPIOS
""", sf_config, expected_types=expected_types)

expected_types = {'Cod. Depto':str,'Cod. Municipio':str, 'CIIU Rev 4 principal':str}
df_base = st_query_to_snowflake_and_return_dataframe("""
    SELECT * FROM TABLA_TEJIDO_MUNICIPIOS
""", sf_config, expected_types=expected_types)

expected_types = {'Código .1':str}
df_ubicacion = st_query_to_snowflake_and_return_dataframe("""
    SELECT * FROM TABLA_DIVIPOLA_MUNICIPIOS
""", sf_config, expected_types=expected_types)

# Devolver los dataframes cargados. Esto es para llamarlos en el principal app.py
def get_dataframes():
    return df_general, df_base, df_ubicacion

def crear_metricas_pdet_zomac(df_datos_mun):
    # Asegurarse de que se está trabajando con una copia del DataFrame para evitar advertencias
    df_datos_mun = df_datos_mun.copy()
    
    # Usar .loc para asignaciones seguras
    df_datos_mun.loc[:, 'PDET'] = 'Es territorio PDET'
    df_datos_mun.loc[:, 'Subregión PDET'] = df_datos_mun['Subregión PDET'].str.title()
    df_datos_mun.loc[:, 'Metrica PDET'] = np.where(df_datos_mun['Subregión PDET'].notna(), df_datos_mun['PDET'] + ' - Subregión ' + df_datos_mun['Subregión PDET'], 'No es territorio PDET')
    df_datos_mun.loc[:, 'Metrica ZOMAC'] = np.where(df_datos_mun['ZOMAC'] == 1, 'Es territorio ZOMAC', 'No es territorio ZOMAC')
    
    return df_datos_mun['Metrica PDET'].values[0], df_datos_mun['Metrica ZOMAC'].values[0]

def mostrar_grafico_torta_datos(df_datos_mun, titulo, etiquetas, valores, colores, texto_central):
    st.subheader(titulo)
    
    fig = go.Figure(data=[go.Pie(labels=etiquetas,
                                 values=valores,
                                 hole=0.5,
                                 marker=dict(colors=colores))])
    
    fig.update_traces(textposition='outside',
                      textinfo='percent+label',
                      hoverinfo='label+percent',
                      textfont=dict(size=16))
    
    fig.update_layout(showlegend=False, annotations=[
        {
            'x': 0.5,
            'y': 0.5,
            'xanchor': 'center',
            'yanchor': 'middle',
            'text': texto_central,
            'showarrow': False,
            'font': {'size': 20}
        }])
    
    st.plotly_chart(fig, use_container_width=True)

def mostrar_empresas_por_categoria_unificada(df, columna_categoria, titulo_seccion, titulo_grafico, color_barras, df_filtrado=None, height=None):
    """
    Muestra información sobre empresas categorizadas por una columna específica, con la opción de utilizar un DataFrame filtrado.

    Args:
        df (pandas.DataFrame): DataFrame que contiene los datos de las empresas.
        columna_categoria (str): Nombre de la columna que se utilizará para categorizar las empresas.
        titulo_seccion (str): Título de la sección que se mostrará en Streamlit.
        titulo_grafico (str): Título del gráfico que se mostrará.
        color_barras (str): Color de las barras en el gráfico.
        df_filtrado (pandas.DataFrame, optional): DataFrame filtrado que contiene las empresas de interés. Por defecto es None.
        height (int, optional): Altura del gráfico en píxeles. Por defecto es None.

    Returns:
        None
    """
    st.subheader(titulo_seccion)
    
    if df_filtrado is not None:
        empresas_filtradas = df_filtrado
    else:
        empresas_filtradas = df
    
    if not empresas_filtradas['Cod. Municipio'].isnull().all():
        c1, c2 = st.columns([20, 80])

        with c1:
            conteo_empresas = pd.pivot_table(empresas_filtradas, index=columna_categoria, values='Número de empresas', aggfunc='sum').reset_index()
            total_empresas = conteo_empresas['Número de empresas'].sum()
            st.markdown('##')
            st.markdown("##### **Cantidad total de empresas**")
            st.subheader(f'{total_empresas:,.0f}')

        with c2:
            conteo_empresas = conteo_empresas.sort_values(by='Número de empresas', ascending=True)
            conteo_empresas['Participación'] = conteo_empresas['Número de empresas'] / conteo_empresas['Número de empresas'].sum() * 100
            text_labels = [f"{num_empresas:,.0f}<br>{participacion:.1f}%"
                           for num_empresas, participacion in zip(conteo_empresas['Número de empresas'], conteo_empresas['Participación'])]

            fig = go.Figure([go.Bar(y=conteo_empresas[columna_categoria],
                                    x=conteo_empresas['Número de empresas'],
                                    text=text_labels,
                                    hoverinfo='text',
                                    orientation='h',
                                    textangle=0,
                                    marker=dict(color=color_barras),
                                    textposition='outside')])
            fig.update_layout(title=titulo_grafico, xaxis_title='Número de empresas', height=height, font=dict(size=16), xaxis=dict(tickfont=dict(size=16)), yaxis=dict(tickfont=dict(size=16)), title_font=dict(size=20))
            st.plotly_chart(fig, use_container_width=True)
    else:
        st.markdown('##')
        st.markdown("##### **Cantidad total de empresas**")
        st.subheader(f'0')

def mostrar_empresas_turismo(tejido):
    st.subheader('Empresas ubicadas en el territorio relacionadas con actividades de turismo')
    turismo = tejido[tejido['Cadena productiva'] == "Turismo"]

    if not turismo['Cod. Municipio'].isnull().all():
        c1, c2 = st.columns([20, 80])

        with c1:
            conteo_empresas6 = pd.pivot_table(turismo, index=['CIIU Rev 4 principal', 'Descripción CIIU principal'], values=['Número de empresas'], aggfunc='sum').reset_index()
            total_tur = conteo_empresas6['Número de empresas'].sum()
            st.markdown('##')
            st.markdown("##### **Cantidad total de empresas**")
            st.subheader(f'{total_tur:,.0f}')

        with c2:
            conteo_empresas6 = conteo_empresas6.sort_values(by='Número de empresas', ascending=True)
            conteo_empresas6['Participación'] = conteo_empresas6['Número de empresas'] / conteo_empresas6['Número de empresas'].sum() * 100
            text_labels = [f"{num_empresas6:,.0f}<br>{participacion6:.1f}%"
                           for num_empresas6, participacion6 in zip(conteo_empresas6['Número de empresas'], conteo_empresas6['Participación'])]

            fig_tur = go.Figure([go.Bar(y=conteo_empresas6['Descripción CIIU principal'],
                                        x=conteo_empresas6['Número de empresas'],
                                        text=text_labels,
                                        hoverinfo='text',
                                        orientation='h',
                                        textangle=0,
                                        marker=dict(color='rgb(255, 218, 0)'),
                                        textposition='outside')])
            fig_tur.update_layout(title='Distribución según CIIU principal', xaxis_title='Número de empresas', height=700, width=800, font=dict(size=16), xaxis=dict(tickfont=dict(size=16)), yaxis=dict(tickfont=dict(size=16)), title_font=dict(size=20))
            st.plotly_chart(fig_tur, use_container_width=True)
    else:
        st.markdown('##')
        st.markdown("##### **Cantidad total de empresas**")
        st.subheader(f'0')    

# FUNCIONES A PROBAR

def generar_reporte_word(municipio, departamento, graficas, texto, cifras):
    # Crear un nuevo documento Word
    doc = docx.Document()

    # Agregar el nombre del municipio y departamento al inicio
    doc.add_heading(f"{municipio} - {departamento}", 0)

    # Agregar el texto
    for parrafo in texto:
        doc.add_paragraph(parrafo)

    # Agregar las cifras
    for cifra in cifras:
        doc.add_paragraph(str(cifra))

    # Agregar las gráficas
    for i, grafica in enumerate(graficas):
        # Guardar temporalmente la gráfica como una imagen
        filename = f"temp_grafica_{i}.png"
        grafica.write_image(filename)
        doc.add_picture(filename)

    # Guardar el documento Word
    doc.save(f"Reporte_{municipio}_{departamento}.docx")

def mostrar_grafico_torta_datos2(df_datos_mun, titulo, etiquetas, valores, colores, texto_central):
    fig = go.Figure(data=[go.Pie(labels=etiquetas,
                                 values=valores,
                                 hole=0.5,
                                 marker=dict(colors=colores))])
    
    fig.update_traces(textposition='outside',
                      textinfo='percent+label',
                      hoverinfo='label+percent',
                      textfont=dict(size=16))
    
    fig.update_layout(title=titulo,
                      annotations=[
                          {
                              'x': 0.5,
                              'y': 0.5,
                              'xanchor': 'center',
                              'yanchor': 'middle',
                              'text': texto_central,
                              'showarrow': False,
                              'font': {'size': 20}
                          }
                      ])
    
    return fig





